import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile

# Caminho para a logo
logo_path = "logo.png"

# Exibir a logo no topo do aplicativo
st.image(logo_path, width=200)
st.title("Gerador de Documentos Personalizados")

st.write("Faça upload de um arquivo de dados (Excel ou CSV) e um modelo de documento Word.")

# Upload do arquivo de dados
data_file = st.file_uploader("Carregar arquivo de dados (Excel ou CSV)", type=["xlsx", "csv"])

# Upload do modelo de documento
template_file = st.file_uploader("Carregar modelo de documento (Word)", type=["docx"])

if data_file and template_file:
    # Ler o arquivo de dados
    if data_file.name.endswith('.xlsx'):
        df = pd.read_excel(data_file)
    else:
        df = pd.read_csv(data_file)

    st.write("Pré-visualização dos dados:")
    st.write(df.head())

    # Mostrar colunas disponíveis para substituição
    st.write("Colunas detectadas no arquivo de dados:", df.columns.tolist())

    # Seleção das colunas para o nome do arquivo
    selected_columns = st.multiselect("Escolha as colunas para o nome do arquivo", df.columns)

    # Botão para gerar documentos personalizados
    if st.button("Gerar documentos personalizados"):
        # Buffer para armazenar o arquivo ZIP
        zip_buffer = io.BytesIO()

        # Criação do arquivo ZIP
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i, row in df.iterrows():
                # Carregar o modelo de documento para cada linha de dados
                doc = Document(template_file)

                # Substituir placeholders no documento
                for paragrafo in doc.paragraphs:
                    for coluna in df.columns:
                        placeholder = f"[{coluna.upper()}]"
                        if placeholder in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(placeholder, str(row[coluna]))

                # Substituir placeholders em tabelas, se houver
                for tabela in doc.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in celula.paragraphs:
                                for coluna in df.columns:
                                    placeholder = f"[{coluna.upper()}]"
                                    if placeholder in paragrafo.text:
                                        paragrafo.text = paragrafo.text.replace(placeholder, str(row[coluna]))

                # Definir o nome do arquivo com base nas colunas selecionadas
                if selected_columns:
                    file_name_parts = [str(row[coluna]) for coluna in selected_columns]
                    file_name = "_".join(file_name_parts) + ".docx"
                else:
                    file_name = f"documento_{i+1}.docx"  # Nome padrão se nenhuma coluna for selecionada

                # Salvar o documento personalizado em um arquivo temporário
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Adicionar o documento ao arquivo ZIP
                zip_file.writestr(file_name, doc_io.read())

        zip_buffer.seek(0)

        # Oferecer o download do arquivo ZIP
        st.download_button(
            label="Baixar todos os documentos em um ZIP",
            data=zip_buffer,
            file_name="documentos_personalizados.zip",
            mime="application/zip"
        )

        st.success("Arquivo ZIP com documentos personalizados gerado com sucesso!")
